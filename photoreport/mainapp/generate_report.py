from docx import Document
from docx.shared import Inches
import os
from django.conf import settings


class DocumentGenerator(object):

    def create(self, project):
        document = Document()

        document.add_heading(project.name, 0)

        for image in project.sorted_images:
            document.add_paragraph(image.caption, style='IntenseQuote')

            img_full_url = os.path.join(
                settings.MEDIA_ROOT, image.image.url.split('/media/')[1]
            )
            document.add_picture(img_full_url, width=Inches(5.25))
            document.add_page_break()

        # document.save(file_name)
        return document
